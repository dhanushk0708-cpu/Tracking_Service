<<<<<<< HEAD
# app.py — updated (drop-in replacement)
import os
from functools import wraps
from urllib.parse import urlparse, urljoin
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re
from flask import send_file
import time
import qrcode
from docx.enum.text import WD_ALIGN_PARAGRAPH


from flask import (
    Flask, render_template, request, jsonify,
    redirect, url_for, session, Response
)
import mysql.connector
import pandas as pd

# ---------------------------------------
# Load config (prefer config.py or environment)
# ---------------------------------------
try:
    from config import (
        MYSQL_HOST, MYSQL_USER, MYSQL_PASSWORD, MYSQL_DB,
        ADMIN_USERNAME, ADMIN_PASSWORD, SECRET_KEY
    )
except Exception:
    MYSQL_HOST = os.environ.get("DB_HOST", "localhost")
    MYSQL_USER = os.environ.get("DB_USER", "root")
    MYSQL_PASSWORD = os.environ.get("DB_PASSWORD", "")
    MYSQL_DB = os.environ.get("DB_NAME", "tracking_service_db")

    ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "admin")
    ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "admin1234")
    SECRET_KEY = os.environ.get("SECRET_KEY", "please-change-me-to-a-random-secret")

DEBUG = os.environ.get("DEBUG", "0") in ("1", "true", "True")

app = Flask(__name__)
app.config["SECRET_KEY"] = SECRET_KEY
app.config["DEBUG"] = DEBUG


# ---------------------------------------
# DB Connection Helper
# ---------------------------------------
def get_db_connection():
    return mysql.connector.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database=MYSQL_DB,
        autocommit=False,
    )


# ---------------------------------------
# Helpers for safety & normalization
# ---------------------------------------
def is_safe_redirect(target):
    if not target:
        return False
    host_url = urlparse(request.host_url)
    redirect_url = urlparse(urljoin(request.host_url, target))
    return (redirect_url.scheme in ("http", "https")
            and host_url.netloc == redirect_url.netloc)


def normalize_value_for_storage(v):
    """
    Normalize phone/pincode when reading from Excel (pandas).
    - If it's NaN -> return empty string
    - If float and integral -> return int string
    - If string like '9908003000.0' -> strip trailing .0
    - Else return stripped string
    """
    if pd.isna(v):
        return ""
    # floats
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        else:
            # keep as is but stringified
            return str(v)
    if isinstance(v, int):
        return str(v)
    s = str(v).strip()
    if s.endswith(".0"):
        s = s[:-2]
    return s


def normalize_value_for_display(v):
    """
    Normalize values before rendering in templates or returning JSON.
    Accepts values from DB (which may already be strings or numbers).
    """
    if v is None:
        return ""
    try:
        # handle bytes, numpy types etc
        # numeric float like 9908003000.0
        if isinstance(v, float):
            if v.is_integer():
                return str(int(v))
            else:
                return str(v)
        if isinstance(v, int):
            return str(v)
        s = str(v)
        if s.endswith(".0"):
            s = s[:-2]
        return s
    except Exception:
        return str(v)







FROM_TEXT = "From:\nSarees by Siva\n+91 9980772558"

def parse_addresses_hybrid(raw_text):
    lines = [l.rstrip() for l in raw_text.splitlines() if l.strip()]

    def is_header(line):
        # WhatsApp Web header
        return bool(re.match(r'^\[.*?\]\s+.*?:', line))

    def strip_header(line):
        """
        Removes FULL WhatsApp header and keeps ONLY address content
        Example:
        [1/22, 3:14 PM] Pavi Akka @ Staff: Name: Bhargavi
        -> Name: Bhargavi
        """
        if "]" in line:
            after = line.split("]", 1)[1]
            if ":" in after:
                return after.split(":", 1)[1].strip()
        return ""

    def is_order(line):
        return line.startswith("Order #")

    addresses = []
    current = []

    has_header = any(is_header(l) for l in lines)

    for line in lines:
        low = line.lower()

        # Remove ONLY allowed junk
        if low.startswith("ship to") or low.startswith("bill to"):
            continue
        if re.match(r'\d{1,2}\s+\w+\s+\d{4}', line):
            continue

        # -------- HEADER MODE (ABSOLUTE PRIORITY) --------
        if has_header and is_header(line):
            if current:
                addresses.append("\n".join(current))
                current = []

            remaining = strip_header(line)
            if remaining:
                current.append(remaining)
            continue

        # Order forces new address
        if is_order(line) and current:
            addresses.append("\n".join(current))
            current = []

        current.append(line)

    if current:
        addresses.append("\n".join(current))

    final_addresses = []

    for addr in addresses:
        addr = addr.strip()
        if not addr:
            continue

        # Add From block ONLY if missing
        if "From:" not in addr and "from:" not in addr:
            addr = addr + "\n\n" + FROM_TEXT + "\n"

        final_addresses.append(addr)

    return final_addresses




def generate_address_word(addresses, output_path):
    doc = Document()

    # -------- Page setup --------
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Narrow margins
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    table = None
    count = 0

    for addr in addresses:
        # New page every 6 addresses (2 × 3)
        if count % 6 == 0:
            if table:
                doc.add_page_break()
            table = doc.add_table(rows=2, cols=3)
            table.style = "Table Grid"  # Visible borders

        row = (count % 6) // 3
        col = (count % 6) % 3
        cell = table.rows[row].cells[col]

        # -------- AUTO FONT SIZE (NO TEXT CHANGE) --------
        length = len(addr)

        if length <= 420:
            font_size = 15
        elif length <= 480:
            font_size = 14
        elif length <= 540:
            font_size = 13
        else:
            font_size = 12

        p = cell.paragraphs[0]
        run = p.add_run(addr)
        run.font.name = "Cambria"
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')

        count += 1

    doc.save(output_path)

# ================= FINAL DETAIL EXTRACTOR =================
def extract_details(text):

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    result = {
        "order_id": "",
        "name": "",
        "pincode": "",
        "phone": ""
    }

    # -------- ORDER ID --------
    order_match = re.search(r"Order\s*#?\s*([A-Z0-9]+)", text, re.IGNORECASE)
    if order_match:
        result["order_id"] = order_match.group(1).upper()

    # -------- PHONE (STRICT INDIAN) --------
    phones = re.findall(r"(?:\+91[\-\s]?|91)?[6-9]\d{9}", text)

    valid_numbers = []
    for ph in phones:
        digits = re.sub(r"\D", "", ph)

        # Remove your FROM number
        if digits.endswith("9980772558"):
            continue

        if len(digits) == 12 and digits.startswith("91"):
            digits = digits[2:]

        if len(digits) == 10:
            valid_numbers.append(digits)

    if valid_numbers:
        result["phone"] = valid_numbers[-1]

    # -------- PINCODE --------
    pin_match = re.search(r"\b\d{6}\b", text)
    if pin_match:
        result["pincode"] = pin_match.group()

    # -------- NAME DETECTION --------

    ignore_words = [
        "india", "state", "dist", "district",
        "road", "street", "colony", "village",
        "mandal", "telangana", "andhra",
        "karnataka", "tamil", "ap", "ts",
        "pincode", "phone", "mobile", "mob"
    ]

    # Case 1: Name prefix
    for line in lines:
        if re.match(r"(?i)^name[\s:\-]", line):
            result["name"] = re.split(r"[:\-]", line, 1)[-1].strip().title()
            return result

    # Case 2: After SHIP TO
    for i, line in enumerate(lines):
        if line.upper() == "SHIP TO" and i + 1 < len(lines):
            if not re.search(r"\d", lines[i + 1]):
                result["name"] = lines[i + 1].title()
                return result

    # Case 3: First clean non-digit short line
    for line in lines:
        lower = line.lower()
        if not re.search(r"\d", line):
            if 1 <= len(line.split()) <= 4 and not any(w in lower for w in ignore_words):
                result["name"] = line.title()
                return result

    return result


# ================= FINAL QR GENERATOR (EXCEL FIXED) =================
def generate_qr_from_address(text, index):

    details = extract_details(text)

    # 🔥 USE TAB → Excel auto splits into next columns
    qr_text = f"{details['name']}\t{details['order_id']}\t{details['pincode']}\t{details['phone']}"

    qr = qrcode.QRCode(
        version=None,
        box_size=3,   # smaller QR
        border=1
    )

    qr.add_data(qr_text)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    temp_dir = os.path.join(BASE_DIR, "temp_qr")
    os.makedirs(temp_dir, exist_ok=True)

    file_name = os.path.join(
        temp_dir,
        f"qr_{int(time.time())}_{index}.png"
    )

    img.save(file_name)

    # Cleanup old files (safe)
    try:
        files = sorted(
            [os.path.join(temp_dir, f) for f in os.listdir(temp_dir)],
            key=os.path.getmtime
        )
        if len(files) > 200:
            for old_file in files[:50]:
                os.remove(old_file)
    except:
        pass

    return file_name


# ================= FINAL WORD GENERATOR (NO BLANK SPACE) =================
def generate_address_word_with_qr(addresses, output_path):

    doc = Document()

    # Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    table = None
    count = 0

    for addr in addresses:

        if count % 6 == 0:
            if table:
                doc.add_page_break()
            table = doc.add_table(rows=2, cols=3)
            table.style = "Table Grid"

        row = (count % 6) // 3
        col = (count % 6) % 3
        cell = table.rows[row].cells[col]

        cell.text = ""

        # Remove cell spacing
        cell.paragraphs[0].paragraph_format.space_before = 0
        cell.paragraphs[0].paragraph_format.space_after = 0

        # QR (Top Right)
        qr_file = generate_qr_from_address(addr, count)

        qr_para = cell.add_paragraph()
        qr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_para.paragraph_format.space_before = 0
        qr_para.paragraph_format.space_after = 2

        qr_run = qr_para.add_run()
        qr_run.add_picture(qr_file, width=Inches(0.5))  # smaller

        # Address below QR
        addr_para = cell.add_paragraph()
        addr_para.paragraph_format.space_before = 0
        addr_para.paragraph_format.space_after = 0

        run = addr_para.add_run(addr)
        run.font.name = "Cambria"
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')

        count += 1

    doc.save(output_path)







# ---------------------------------------
# Basic Auth (Browser popup)
# ---------------------------------------
def check_basic_auth(username, password):
    return username == ADMIN_USERNAME and password == ADMIN_PASSWORD


def basic_auth_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if session.get("admin_logged_in"):
            return view(*args, **kwargs)

        auth = request.authorization
        if auth and check_basic_auth(auth.username, auth.password):
            return view(*args, **kwargs)

        return Response(
            "Authentication required", 401,
            {"WWW-Authenticate": 'Basic realm="Admin Area"'}
        )

    return wrapped


# ---------------------------------------
# HTML Login (Session)
# ---------------------------------------
def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if session.get("admin_locked_in") or session.get("admin_logged_in"):
            return view(*args, **kwargs)
        next_url = request.path
        return redirect(url_for("admin_login", next=next_url))
    return wrapped


# ---------------------------------------
# ROUTES
# ---------------------------------------
=======
import os
from flask import (
    Flask,
    render_template,
    request,
    jsonify,
    redirect,
    url_for,
    session,
)
import mysql.connector
import pandas as pd
from functools import wraps

app = Flask(__name__)

# -----------------------------------------
# SECRET KEY (REQUIRED FOR LOGIN SESSIONS)
# -----------------------------------------
# -----------------------------------------
# SECRET KEY (REQUIRED FOR LOGIN SESSIONS)
# -----------------------------------------
app.secret_key = os.environ.get("SECRET_KEY", "2507")

# -----------------------------------------
# SIMPLE ADMIN CREDENTIALS (OVERRIDDEN BY ENV VARS IN PRODUCTION)
# -----------------------------------------
ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "admin1234")

# -----------------------------------------
# DATABASE CONFIG
# -----------------------------------------
DB_CONFIG = {
    "host": os.environ.get("DB_HOST", "localhost"),
    "user": os.environ.get("DB_USER", "root"),
    "password": os.environ.get("DB_PASSWORD", "2507"),
    "database": os.environ.get("DB_NAME", "tracking_service_db")
}


def get_db_connection():
    return mysql.connector.connect(**DB_CONFIG)


def norm(name: str) -> str:
    """Normalize Excel column names."""
    return str(name).strip().lower().replace(" ", "").replace(".", "")


# -----------------------------------------
# LOGIN REQUIRED DECORATOR
# -----------------------------------------
def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("admin_logged_in"):
            next_url = request.path
            return redirect(url_for("admin_login", next=next_url))
        return f(*args, **kwargs)
    return wrapper


# -----------------------------------------
# HOME ROUTE
# -----------------------------------------
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
@app.route("/")
def home():
    return "Tracking Service is Running!"


<<<<<<< HEAD
=======
# -----------------------------------------
# TEST DB
# -----------------------------------------
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
@app.route("/test-db")
def test_db():
    try:
        conn = get_db_connection()
        conn.close()
        return "MySQL Connection Successful!"
    except Exception as e:
<<<<<<< HEAD
        return f"MySQL Connection Failed: {e}", 500


# ---------------------------------------
# ADMIN LOGIN PAGE
# ---------------------------------------
=======
        return f"MySQL Connection Failed: {e}"


# -----------------------------------------
# ADMIN LOGIN / LOGOUT
# -----------------------------------------
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    error = None
    next_url = request.args.get("next") or url_for("admin_dashboard")

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
<<<<<<< HEAD

        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session["admin_logged_in"] = True
            if is_safe_redirect(request.form.get("next")):
                return redirect(request.form.get("next"))
            return redirect(url_for("admin_dashboard"))

        error = "Invalid username or password."
=======
        next_url = request.form.get("next") or url_for("admin_dashboard")

        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session["admin_logged_in"] = True
            return redirect(next_url)
        else:
            error = "Invalid username or password."
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad

    return render_template("login.html", error=error, next=next_url)


@app.route("/admin/logout")
def admin_logout():
    session.clear()
    return redirect(url_for("admin_login"))


<<<<<<< HEAD
# ---------------------------------------
# ADMIN DASHBOARD
# ---------------------------------------
@app.route("/admin")
@basic_auth_required
=======
# -----------------------------------------
# ADMIN: UPLOAD EXCEL (PROTECTED)
# -----------------------------------------
@app.route("/admin/upload", methods=["GET", "POST"])
@login_required
def upload_page():
    if request.method == "GET":
        return render_template("upload.html")

    file = request.files.get("file")
    if not file:
        return "No file selected"

    try:
        df = pd.read_excel(file)
        df.columns = [norm(c) for c in df.columns]

        required_cols = [
            "slno",          # Sl.no
            "name",
            "orderid",
            "pincode",
            "phonenumber",
            "tracknumber",
            "weight",
            "couriername",
        ]
        for col in required_cols:
            if col not in df.columns:
                return f"Missing column in Excel (after normalize): {col}"

        conn = get_db_connection()
        cursor = conn.cursor()
        inserted = 0

        for _, row in df.iterrows():
            customer_name = str(row["name"]).strip()
            order_id = str(row["orderid"]).strip()
            phone = str(row["phonenumber"]).strip()
            pincode = str(row["pincode"]).strip()
            tracking_number = str(row["tracknumber"]).strip()
            weight = str(row["weight"]).strip()
            courier_name = str(row["couriername"]).strip()

            courier_name_lower = courier_name.lower()
            if courier_name_lower == "dtdc":
                courier_site = (
                    "https://www.dtdc.in/tracking.asp?"
                    f"Ttype=awb&strCNNo={tracking_number}"
                )
            elif courier_name_lower in ["india post", "indian post"]:
                courier_site = (
                    "https://www.indiapost.gov.in/_layouts/15/"
                    f"dop.portal.tracking/trackconsignment.aspx?consignmentno={tracking_number}"
                )
            else:
                courier_site = ""

            sql = """
                INSERT INTO shipments
                (customer_name, order_id, phone, pincode,
                 tracking_number, weight, courier_name, courier_site)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """
            values = (
                customer_name,
                order_id,
                phone,
                pincode,
                tracking_number,
                weight,
                courier_name,
                courier_site,
            )
            cursor.execute(sql, values)
            inserted += 1

        conn.commit()
        cursor.close()
        conn.close()

        return f"Excel uploaded! Total rows inserted: {inserted}"

    except Exception as e:
        return f"Error processing file: {e}"


# -----------------------------------------
# PUBLIC API: TRACK BY PHONE OR ORDER ID
# -----------------------------------------
@app.route("/api/track", methods=["GET"])
def track_by_phone_or_order():
    phone = request.args.get("phone", "").strip()
    order_id = request.args.get("order_id", "").strip()

    if not phone and not order_id:
        return jsonify({"error": "Provide phone or order_id parameter"}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        where_clauses = []
        params = []

        if phone:
            where_clauses.append("phone = %s")
            params.append(phone)

        if order_id:
            where_clauses.append("order_id = %s")
            params.append(order_id)

        where_sql = " AND ".join(where_clauses)

        query = f"""
            SELECT customer_name, order_id, phone, pincode,
                   tracking_number, weight, courier_name, courier_site, updated_at
            FROM shipments
            WHERE {where_sql}
            ORDER BY updated_at DESC
        """

        cursor.execute(query, params)
        rows = cursor.fetchall()
        cursor.close()
        conn.close()

        if not rows:
            return jsonify({"results": [], "message": "No orders found for given details."})

        return jsonify({"results": rows})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------------------------
# ADMIN DASHBOARD (PROTECTED)
# -----------------------------------------
@app.route("/admin")
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
@login_required
def admin_dashboard():
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")

<<<<<<< HEAD
    conn = None
    cursor = None
=======
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT COUNT(*) FROM shipments")
<<<<<<< HEAD
        overall_total = cursor.fetchone()[0] or 0

        where = []
        params = []

        if start_date:
            where.append("DATE(updated_at) >= %s")
            params.append(start_date)

        if end_date:
            where.append("DATE(updated_at) <= %s")
            params.append(end_date)

        where_sql = " WHERE " + " AND ".join(where) if where else ""

        cursor.execute("SELECT COUNT(*) FROM shipments" + where_sql, params)
        total_orders = cursor.fetchone()[0] or 0

        cursor.execute(
            "SELECT courier_name, COUNT(*) FROM shipments "
            + where_sql +
            " GROUP BY courier_name",
            params
        )
        courier_stats = cursor.fetchall()

=======
        overall_total = cursor.fetchone()[0]

        where_clauses = []
        params = []

        if start_date:
            where_clauses.append("DATE(updated_at) >= %s")
            params.append(start_date)

        if end_date:
            where_clauses.append("DATE(updated_at) <= %s")
            params.append(end_date)

        where_sql = ""
        if where_clauses:
            where_sql = " WHERE " + " AND ".join(where_clauses)

        cursor.execute("SELECT COUNT(*) FROM shipments" + where_sql, params)
        total_orders = cursor.fetchone()[0]

        courier_query = """
            SELECT courier_name, COUNT(*)
            FROM shipments
        """ + where_sql + " GROUP BY courier_name"

        cursor.execute(courier_query, params)
        courier_stats = cursor.fetchall()

        cursor.close()
        conn.close()

>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
        return render_template(
            "admin.html",
            overall_total=overall_total,
            total_orders=total_orders,
            courier_stats=courier_stats,
            start_date=start_date,
            end_date=end_date,
        )
<<<<<<< HEAD

    except Exception as e:
        return f"Error loading admin dashboard: {e}", 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


# -------------------------
# ADMIN ORDERS PAGE
# -------------------------
@app.route("/admin/orders")
@basic_auth_required
@login_required
def admin_orders():
    query = request.args.get("q", "").strip()
    start_date = request.args.get("start_date", "").strip()
    end_date = request.args.get("end_date", "").strip()
    courier_filter = request.args.get("courier", "").strip()

    conn = None
    cursor = None
=======
    except Exception as e:
        return f"Error loading admin dashboard: {e}"


# -----------------------------------------
# ADMIN ORDER HISTORY (PROTECTED)
# -----------------------------------------
@app.route("/admin/orders")
@login_required
def admin_orders():
    query = request.args.get("q", "").strip()
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    courier_filter = request.args.get("courier", "").strip()

>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

<<<<<<< HEAD
        sql = "SELECT * FROM shipments WHERE 1=1"
        params = []

        if query:
            sql += " AND (customer_name LIKE %s OR phone LIKE %s OR order_id LIKE %s)"
            params += [f"%{query}%", f"%{query}%", f"%{query}%"]

        if start_date:
            sql += " AND DATE(updated_at) >= %s"
            params.append(start_date)

        if end_date:
            sql += " AND DATE(updated_at) <= %s"
            params.append(end_date)

        if courier_filter:
            sql += " AND courier_name = %s"
            params.append(courier_filter)

        sql += " ORDER BY updated_at DESC"
        cursor.execute(sql, params)
        orders = cursor.fetchall() or []

        # normalize phone/pincode/updated_at for display
        for r in orders:
            r["phone"] = normalize_value_for_display(r.get("phone"))
            r["pincode"] = normalize_value_for_display(r.get("pincode"))
            # updated_at to string
            if r.get("updated_at") is not None:
                r["updated_at"] = str(r.get("updated_at"))

        # courier list
        cursor.execute("SELECT DISTINCT courier_name FROM shipments")
        couriers = [row["courier_name"] for row in cursor.fetchall() if row.get("courier_name")]
=======
        where_clauses = []
        params = []

        if query:
            where_clauses.append("""
                (
                    customer_name LIKE %s OR 
                    phone LIKE %s OR 
                    order_id LIKE %s OR 
                    courier_name LIKE %s
                )
            """)
            wildcard = f"%{query}%"
            params.extend([wildcard, wildcard, wildcard, wildcard])

        if start_date:
            where_clauses.append("DATE(updated_at) >= %s")
            params.append(start_date)

        if end_date:
            where_clauses.append("DATE(updated_at) <= %s")
            params.append(end_date)

        if courier_filter:
            where_clauses.append("courier_name = %s")
            params.append(courier_filter)

        where_sql = ""
        if where_clauses:
            where_sql = " WHERE " + " AND ".join(where_clauses)

        query_sql = f"""
            SELECT *
            FROM shipments
            {where_sql}
            ORDER BY updated_at DESC
        """

        cursor.execute(query_sql, params)
        orders = cursor.fetchall()

        cursor.execute("SELECT DISTINCT courier_name FROM shipments")
        courier_rows = cursor.fetchall()
        couriers = [row["courier_name"] for row in courier_rows if row["courier_name"]]

        cursor.close()
        conn.close()
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad

        return render_template(
            "admin_orders.html",
            orders=orders,
            query=query,
            start_date=start_date,
            end_date=end_date,
            couriers=couriers,
            courier_filter=courier_filter
        )

    except Exception as e:
<<<<<<< HEAD
        return f"Error loading orders: {e}", 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()





# -------------------------
# Admin: edit single order (GET form + POST update)
# -------------------------
@app.route("/admin/order/<int:order_id>/edit", methods=["GET", "POST"])
@basic_auth_required
@login_required
def admin_edit_order(order_id):
    conn = None
    cursor = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        if request.method == "POST":
            customer_name = request.form.get("customer_name", "").strip()
            order_id_field = request.form.get("order_id_field", "").strip()
            phone = request.form.get("phone", "").strip()
            pincode = request.form.get("pincode", "").strip()
            tracking_number = request.form.get("tracking_number", "").strip()
            courier_name = request.form.get("courier_name", "").strip()
            courier_site = request.form.get("courier_site", "").strip() or ""

            # allow optional updated_at from admin edit if posted (string)
            updated_at_val = request.form.get("updated_at")
            if updated_at_val:
                update_sql = """
                    UPDATE shipments
                    SET customer_name=%s, order_id=%s, phone=%s, pincode=%s,
                        tracking_number=%s, courier_name=%s, courier_site=%s, updated_at=%s
                    WHERE id = %s
                """
                params = (
                    customer_name, order_id_field, phone, pincode,
                    tracking_number, courier_name, courier_site, updated_at_val, order_id
                )
            else:
                update_sql = """
                    UPDATE shipments
                    SET customer_name=%s, order_id=%s, phone=%s, pincode=%s,
                        tracking_number=%s, courier_name=%s, courier_site=%s
                    WHERE id = %s
                """
                params = (
                    customer_name, order_id_field, phone, pincode,
                    tracking_number, courier_name, courier_site, order_id
                )

            cursor.execute(update_sql, params)
            conn.commit()
            return redirect(url_for("admin_orders"))

        cursor.execute("SELECT * FROM shipments WHERE id = %s", (order_id,))
        order = cursor.fetchone()

        if not order:
            return "Order not found", 404

        # normalize values for form fields
        order["phone"] = normalize_value_for_display(order.get("phone"))
        order["pincode"] = normalize_value_for_display(order.get("pincode"))
        if order.get("updated_at") is not None:
            order["updated_at"] = str(order.get("updated_at"))

        return render_template("admin_edit.html", order=order)

    except Exception as e:
        return f"Error editing order: {e}", 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()



@app.route("/admin/address-tool", methods=["GET", "POST"])
@basic_auth_required
@login_required
def admin_address_tool():
    preview = []

    if request.method == "POST":
        raw_text = request.form.get("raw_text", "").strip()
        action = request.form.get("action")

        parsed_addresses = parse_addresses_hybrid(raw_text)

        # -------- PREVIEW --------
        if action == "preview":
            preview = parsed_addresses

        # -------- GENERATE WORD --------
        elif action == "generate":
            if not parsed_addresses:
                return "No addresses found", 400

            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            output_dir = os.path.join(BASE_DIR, "output")
            os.makedirs(output_dir, exist_ok=True)

            filename = f"addresses_{int(time.time())}.docx"
            output_file = os.path.join(output_dir, filename)

            generate_address_word(parsed_addresses, output_file)

            return send_file(
                output_file,
                as_attachment=True,
                download_name=filename
            )

    return render_template("admin_address_tool.html", preview=preview)



@app.route("/admin/address-tool-qr", methods=["GET", "POST"])
@basic_auth_required
@login_required
def admin_address_tool_qr():
    preview = []

    if request.method == "POST":
        raw_text = request.form.get("raw_text", "").strip()
        action = request.form.get("action")

        parsed_addresses = parse_addresses_hybrid(raw_text)

        if action == "preview":
            preview = parsed_addresses

        elif action == "generate":
            if not parsed_addresses:
                return "No addresses found", 400

            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            output_dir = os.path.join(BASE_DIR, "output")
            os.makedirs(output_dir, exist_ok=True)

            filename = f"addresses_qr_{int(time.time())}.docx"
            output_file = os.path.join(output_dir, filename)

            generate_address_word_with_qr(parsed_addresses, output_file)

            return send_file(
                output_file,
                as_attachment=True,
                download_name=filename
            )

    return render_template("admin_address_tool.html", preview=preview)



# -------------------------
# Admin: bulk edit (simple actions)
# Accepts POST to /admin/orders/bulk_edit
# -------------------------
@app.route("/admin/orders/bulk_edit", methods=["POST"])
@app.route("/admin/bulk_edit", methods=["POST"])  # alias for older template forms
@basic_auth_required
@login_required
def admin_orders_bulk_edit():
    try:
        # accept both multiple selected_ids inputs and a single comma-separated string
        ids = request.form.getlist("selected_ids") or []
        if len(ids) == 1 and ids[0] and ',' in ids[0]:
            ids = [s.strip() for s in ids[0].split(',') if s.strip()]

        # also accept front-end 'action' param
        action = request.form.get("action") or request.form.get("bulk_action")
        value = request.form.get("bulk_value", "").strip()

        if not ids:
            return "No rows selected", 400

        # convert to ints
        try:
            ids_int = [int(i) for i in ids]
        except Exception:
            return "Invalid id list", 400

        conn = get_db_connection()
        cursor = conn.cursor()

        placeholders = ",".join(["%s"] * len(ids_int))

        if action == "set-courier" or request.form.get("bulk_action") == "set-courier":
            sql = f"UPDATE shipments SET courier_name = %s WHERE id IN ({placeholders})"
            params = [value] + ids_int
            cursor.execute(sql, params)

        elif action == "clear-courier" or request.form.get("bulk_action") == "clear-courier":
            sql = f"UPDATE shipments SET courier_name = '' WHERE id IN ({placeholders})"
            cursor.execute(sql, ids_int)

        elif action == "delete" or request.form.get("action") == "delete" or request.form.get("bulk_action") == "delete":
            sql = f"DELETE FROM shipments WHERE id IN ({placeholders})"
            cursor.execute(sql, ids_int)

        else:
            cursor.close()
            conn.close()
            return "Unknown action", 400

        conn.commit()
        cursor.close()
        conn.close()
        return redirect(url_for("admin_orders"))

    except Exception as e:
        return f"Error in bulk edit: {e}", 500


# -------------------------
# Admin: bulk edit via upload (CSV / Excel)
# -------------------------
@app.route("/admin/bulk-edit", methods=["GET", "POST"])
@basic_auth_required
@login_required
def admin_bulk_edit():
    if request.method == "GET":
        return render_template("admin_bulk_edit.html")

    file = request.files.get("file")
    if not file:
        return "No file uploaded", 400

    try:
        import io
        filename = file.filename.lower()
        file.stream.seek(0)
        if filename.endswith(".csv"):
            df = pd.read_csv(io.StringIO(file.read().decode("utf-8")))
        else:
            file.stream.seek(0)
            df = pd.read_excel(file)
    except Exception as e:
        return f"Error reading uploaded file: {e}", 400

    df.columns = [str(c).strip().lower().replace(" ", "").replace(".", "") for c in df.columns]

    allowed = {
        "customer_name": "customer_name",
        "name": "customer_name",
        "orderid": "order_id",
        "order_id": "order_id",
        "phonenumber": "phone",
        "phone": "phone",
        "pincode": "pincode",
        "tracknumber": "tracking_number",
        "tracking_number": "tracking_number",
        "trackingnumber": "tracking_number",
        "couriername": "courier_name",
        "courier_name": "courier_name",
        "updated_at": "updated_at",
        "id": "id"
    }

    updated = 0
    errors = []
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        for idx, row in df.iterrows():
            target_id = None
            if "id" in df.columns and not pd.isna(row.get("id")):
                try:
                    target_id = int(row.get("id"))
                except Exception:
                    target_id = None

            if not target_id:
                tnum = None
                for colname in ("tracknumber", "tracking_number", "trackingnumber"):
                    if colname in df.columns and not pd.isna(row.get(colname)):
                        tnum = str(row.get(colname)).strip()
                        break
                if tnum:
                    cursor.execute("SELECT id FROM shipments WHERE tracking_number = %s LIMIT 1", (tnum,))
                    r = cursor.fetchone()
                    if r:
                        target_id = r[0]

            if not target_id:
                errors.append(f"Row {idx+1}: no id or tracking_number to match")
                continue

            updates = []
            params = []
            for col_in, col_db in allowed.items():
                if col_in in ("id", "tracknumber", "trackingnumber", "tracking_number"):
                    continue
                if col_in in df.columns and not pd.isna(row.get(col_in)):
                    val = row.get(col_in)
                    if col_db in ("phone", "pincode"):
                        val = normalize_value_for_storage(val)
                    else:
                        # convert pandas timestamps where appropriate
                        if col_db == "updated_at":
                            # normalized later
                            pass
                    updates.append(f"{col_db} = %s")
                    params.append(val if not pd.isna(val) else None)

            if not updates:
                continue

            params.append(target_id)
            sql = "UPDATE shipments SET " + ", ".join(updates) + " WHERE id = %s"
            try:
                cursor.execute(sql, tuple(params))
                updated += cursor.rowcount
            except Exception as e:
                errors.append(f"Row {idx+1} update error: {e}")

        conn.commit()
    finally:
        cursor.close()
        conn.close()

    return jsonify({"updated": updated, "errors": errors})


# ---------------------------------------
# ADMIN UPLOAD EXCEL (single upload)
# ---------------------------------------
@app.route("/admin/upload", methods=["GET", "POST"])
@basic_auth_required
@login_required
def upload_page():
    if request.method == "GET":
        return render_template("upload.html")

    file = request.files.get("file")
    if not file:
        return "No file selected", 400

    try:
        file.stream.seek(0)
        df = pd.read_excel(file)
    except Exception as e:
        return f"Error reading Excel file: {e}", 400

    df.columns = [c.strip().lower().replace(" ", "").replace(".", "") for c in df.columns]

    required_cols = [
        "slno", "name", "orderid", "pincode",
        "phonenumber", "tracknumber", "weight", "couriername"
    ]

    for col in required_cols:
        if col not in df.columns:
            return f"Missing column in Excel: {col}", 400

    has_updated_at = "updated_at" in df.columns

    conn = None
    cursor = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        inserted = 0

        for _, row in df.iterrows():
            customer_name = str(row["name"]).strip()
            order_id = str(row["orderid"]).strip()
            phone = normalize_value_for_storage(row["phonenumber"])
            pincode = normalize_value_for_storage(row["pincode"])
            tracking_number = str(row["tracknumber"]).strip()
            weight = str(row["weight"]).strip()
            courier_name = str(row["couriername"]).strip()

            courier_name_lower = courier_name.lower()

            if "dtdc" in courier_name_lower:
                courier_site = "https://www.dtdc.com/track-your-shipment/"
            elif courier_name_lower in ["india post", "indian post"]:
                courier_site = "https://www.indiapost.gov.in/"


            else:
                courier_site = ""

            if has_updated_at:
                updated_at_val = row.get("updated_at")
                try:
                    if pd.isna(updated_at_val):
                        updated_at_val = None
                    else:
                        if hasattr(updated_at_val, "to_pydatetime"):
                            updated_at_val = updated_at_val.to_pydatetime()
                except Exception:
                    updated_at_val = str(updated_at_val)

                sql = """
                    INSERT INTO shipments
                    (customer_name, order_id, phone, pincode,
                     tracking_number, weight, courier_name, courier_site, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                cursor.execute(sql, (
                    customer_name, order_id, phone, pincode,
                    tracking_number, weight, courier_name, courier_site,
                    updated_at_val
                ))
            else:
                sql = """
                    INSERT INTO shipments
                    (customer_name, order_id, phone, pincode,
                    tracking_number, weight, courier_name, courier_site)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """
                cursor.execute(sql, (
                    customer_name, order_id, phone, pincode,
                    tracking_number, weight, courier_name, courier_site
                ))

            inserted += 1

        conn.commit()
        return f"Excel uploaded! Total rows inserted: {inserted}"

    except Exception as e:
        if conn:
            conn.rollback()
        return f"Error processing file: {e}", 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


# ---------------------------------------
# PUBLIC API
# ---------------------------------------
@app.route("/api/track")
def track_api():
    phone = request.args.get("phone", "").strip()
    order_id = request.args.get("order_id", "").strip()

    if not phone and not order_id:
        return jsonify({"error": "Provide phone or order_id"}), 400

    conn = None
    cursor = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        where = []
        params = []

        if phone:
            where.append("phone = %s")
            params.append(phone)
        if order_id:
            where.append("order_id = %s")
            params.append(order_id)

        where_sql = " AND ".join(where)


        cursor.execute(f"""
            SELECT * FROM shipments
            WHERE {where_sql}
            AND updated_at >= NOW() - INTERVAL 15 DAY
            ORDER BY updated_at DESC
            """, params)

        rows = cursor.fetchall() or []

        # normalize numeric values for JSON
        for r in rows:
            r["phone"] = normalize_value_for_display(r.get("phone"))
            r["pincode"] = normalize_value_for_display(r.get("pincode"))
            if r.get("updated_at") is not None:
                r["updated_at"] = str(r.get("updated_at"))

        return jsonify({"results": rows})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


# ---------------------------------------
# PUBLIC TRACK PAGE
# ---------------------------------------
=======
        return f"Error loading order history: {e}"


# -----------------------------------------
# PUBLIC TRACK PAGE (HTML)
# -----------------------------------------
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
@app.route("/track")
def track_page():
    return render_template("track.html")


<<<<<<< HEAD
# ---------------------------------------
# MAIN (only for local/dev)
# ---------------------------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=DEBUG)
=======
# -----------------------------------------
# START SERVER
# -----------------------------------------
if __name__ == "__main__":
    app.run(debug=True)
>>>>>>> 016b26d6750c0d8606ad5b7e7a6fe302a54e72ad
